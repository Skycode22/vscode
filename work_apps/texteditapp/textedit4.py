import os
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.popup import Popup
from docx import Document
from kivy.config import Config
Config.set('graphics', 'width', '1080')
Config.set('graphics', 'height', '1920')

class MainApp(App):
    def build(self):
        main_layout = BoxLayout(orientation="vertical", padding=1, spacing=5)
        self.text_input = TextInput(hint_text="Type your text here", multiline=True, size_hint_y=0.8, cursor_color=(0, 1, 0, 1))
        main_layout.add_widget(self.text_input)

        button_layout = BoxLayout(size_hint_y=0.1)
        open_button = Button(text="Open", background_color=(.75, 1, .79, 1))
        open_button.bind(on_press=self.open_file)
        button_layout.add_widget(open_button)

        save_button = Button(text="Save", background_color=(.75, .5, .79, 1))
        save_button.bind(on_press=self.save_file)
        button_layout.add_widget(save_button)

        main_layout.add_widget(button_layout)
        self.current_file = None
        return main_layout

    def open_file(self, _):
        file_chooser = FileChooserIconView(path=os.getcwd(), filters=["*.docx", "*.txt", "*"])
        file_chooser.bind(on_submit=self.load_file)
        popup = Popup(title="Open text file", content=file_chooser, size_hint=(0.9, 0.9))
        self.file_chooser_popup = popup
        self.file_chooser_popup.open()

    def load_file(self, _, selected, *args):
        if not selected:
            return

        self.current_file = selected[0]
        if self.current_file.endswith('.docx'):
            document = Document(self.current_file)
            text = "\n\n".join([para.text for para in document.paragraphs])
        else:
            with open(self.current_file, 'r', encoding='utf-8') as f:
                text = f.read()

        self.text_input.text = text
        self.file_chooser_popup.dismiss()

    def save_file(self, _):
        if self.current_file is not None:
            self.save_current_file()
        else:
            file_chooser = FileChooserIconView(path=os.getcwd(), filters=["*.docx", "*.txt", "*"])
            file_chooser.bind(on_submit=self.save_new_file)
            popup = Popup(title="Save text file", content=file_chooser, size_hint=(0.9, 0.9))
            self.file_chooser_popup = popup
            self.file_chooser_popup.open()

    def save_current_file(self):
        if self.current_file.endswith('.docx'):
            document = Document()
            text = self.text_input.text
            paragraphs = text.split("\n\n")
            for para in paragraphs:
                document.add_paragraph(para)

            document.save(self.current_file)
        else:
            with open(self.current_file, 'w', encoding='utf-8') as f:
                f.write(self.text_input.text)

    def save_new_file(self, _, selected, *args):
        if not selected:
            return

        self.current_file = selected[0]
        self.save_current_file()
        self.file_chooser_popup.dismiss()

if __name__ == "__main__":
    MainApp().run()
