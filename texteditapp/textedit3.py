import os
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.label import Label
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.textinput import TextInput
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

class FileChooserScreen(Screen):
    def __init__(self, **kwargs):
        super(FileChooserScreen, self).__init__(**kwargs)
        self.layout = BoxLayout(orientation='vertical')
        self.add_widget(self.layout)

        self.file_chooser = FileChooserIconView(path=os.path.expanduser('~'), filters=['*.docx', '*.xlsx'])
        self.file_chooser.bind(on_submit=self.load_document)
        self.layout.add_widget(self.file_chooser)

    def load_document(self, instance, selected, touch):
        if not selected:
            return
        self.manager.get_screen('editor').load_document(selected[0])
        self.manager.current = 'editor'

class EditorScreen(Screen):
    def __init__(self, **kwargs):
        super(EditorScreen, self).__init__(**kwargs)
        self.layout = BoxLayout(orientation='vertical')
        self.add_widget(self.layout)

        self.text_input = TextInput()
        self.layout.add_widget(self.text_input)

        button_layout = BoxLayout(size_hint_y=None, height=50)
        self.layout.add_widget(button_layout)

        self.save_button = Button(text='Save')
        self.save_button.bind(on_press=self.save_document)
        button_layout.add_widget(self.save_button)

        self.open_button = Button(text='Open')
        self.open_button.bind(on_press=self.open_file)
        button_layout.add_widget(self.open_button)

    def load_document(self, file_path):
        self.file_path = file_path
        file_extension = os.path.splitext(file_path)[1]

        if file_extension == '.docx':
            self.file_type = 'docx'
            self.doc = Document(file_path)
            self.text_input.text = '\n'.join([para.text for para in self.doc.paragraphs])
        elif file_extension == '.xlsx':
            self.file_type = 'xlsx'
            self.wb = load_workbook(file_path)
            self.ws = self.wb.active
            rows = self.ws.iter_rows()
            self.text_input.text = '\n'.join([','.join([str(cell.value) for cell in row]) for row in rows])

    def save_document(self, instance):
            if not hasattr(self, 'file_type'):
                return

            text = self.text_input.text
            lines = text.split('\n')

            if self.file_type == 'docx':
                for index, line in enumerate(lines):
                    if index < len(self.doc.paragraphs):
                        self.doc.paragraphs[index].text = line
                    else:
                        self.doc.add_paragraph(line)
                self.doc.save(self.file_path)
            elif self.file_type == 'xlsx':
                for row_index, line in enumerate(lines):
                    cells = line.split(',')
                    for col_index, cell_value in enumerate(cells):
                        self.ws.cell(row=row_index + 1, column=col_index + 1, value=cell_value)
                self.wb.save(self.file_path)

    def open_file(self, instance):
            self.manager.current = 'filechooser'

class DocxEditorApp(App):
    def build(self):
        screen_manager = ScreenManager()
        screen_manager.add_widget(FileChooserScreen(name='filechooser'))
        screen_manager.add_widget(EditorScreen(name='editor'))
        return screen_manager

if __name__ == '__main__':
    DocxEditorApp().run()

