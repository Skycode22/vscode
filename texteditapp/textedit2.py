import os
import docx
import openpyxl
import PyPDF2
from kivy.lang import Builder
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.popup import Popup
from kivy.properties import ObjectProperty
from kivy.app import App
from kivy.uix.button import Button
from kivy.uix.filechooser import FileChooserIconView
from kivy.core.window import Window

Window.size = (360, 640)
Window.clearcolor = (.5, 1, 1, 1)
KV = '''
BoxLayout:
    orientation: "vertical"
    BoxLayout:
        size_hint_y: None
        height: "48dp"
        Button:
            text: "Open"
            on_release: app.open_file()
        Button:
            text: "Save"
            on_release: app.save_file()
        Button:
            text: "Save As"
            on_release: app.save_file_as()        
    ScrollView:
        do_scroll_x: False
        bar_width: "10dp"
        scroll_type: ['bars']
        effect_cls: "ScrollEffect"
        TextInput:
            id: text_input
            focus: True
            font_size: '16sp'
            size_hint_y: None
            height: max(self.minimum_height, root.height - 48)
            on_parent: app.text_input = self
'''
class TextEditor(App):
    text_input = ObjectProperty(None)
    current_file = ''
    def build(self):
        return Builder.load_string(KV)
    def open_file(self):
        filechooser = FileChooserIconView(path=os.getcwd())
        popup = Popup(title="Open File", content=filechooser, size_hint=(0.8, 0.8))
        filechooser.bind(on_submit=self.load_file)
        popup.open()
    def load_file(self, instance, filepaths, *args):
        if filepaths:
            self.current_file = filepaths[0]
            file_ext = os.path.splitext(self.current_file)[1].lower()
            if file_ext == ".docx":
                try:
                    doc = docx.Document(self.current_file)
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text is not None:
                            full_text.append(paragraph.text)
                    self.text_input.text = "\n".join(full_text)
                except Exception as e:
                    print(f"Error: {e}")
            elif file_ext == ".xlsx":
                try:
                    wb = openpyxl.load_workbook(self.current_file)
                    ws = wb.active
                    full_text = []
                    for row in ws.iter_rows():
                        row_text = []
                        for cell in row:
                            row_text.append(str(cell.value))
                        full_text.append('\t'.join(row_text))
                    self.text_input.text = "\n".join(full_text)
                except Exception as e:
                    print(f"Error: {e}")
            elif file_ext == ".pdf":
                try:
                    with open(self.current_file, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        full_text = []
                        for page in pdf_reader.pages:
                            full_text.append(page.extract_text())
                        self.text_input.text = "\n".join(full_text)
                except Exception as e:
                    print(f"Error: {e}")
            else:
                try:
                    with open(self.current_file, 'r', encoding='utf-8') as f:
                        self.text_input.text = f.read()
                except UnicodeDecodeError as e:
                    print(f"Error: {e}")
            instance.parent.parent.parent.dismiss()
    def save_file(self):
        if self.current_file:
            file_ext = os.path.splitext(self.current_file)[1].lower()
            if file_ext == ".docx":
                doc = docx.Document()
                for paragraph_text in self.text_input.text.split("\n"):
                    doc.add_paragraph(paragraph_text)
                doc.save(self.current_file)
            elif file_ext == ".xlsx":
                wb = openpyxl.Workbook()
                ws = wb.active
                for row in self.text_input.text.split("\n"):
                    row_cells = row.split("\t")
                    ws.append(row_cells)
                wb.save(self.current_file)
            else:
                with open(self.current_file, 'w', encoding='utf-8') as f:
                    f.write(self.text_input.text)
        else:
            self.save_file_as()
    def save_file_as(self):
        filechooser = FileChooserIconView(path=os.getcwd(), dirselect=True)
        popup = Popup(title="Save File As", content=filechooser, size_hint=(0.8, 0.8))
        filechooser.bind(on_submit=self.save_new_file_with_popup)
        popup.open()
    def save_new_file_with_popup(self, instance, filepaths, *args):
        if filepaths:
            filename = os.path.basename(instance.selection[0])
            self.current_file = os.path.join(filepaths[0], filename)
            file_ext = os.path.splitext(self.current_file)[1].lower()
            if file_ext == ".docx":
                doc = docx.Document()
                for paragraph_text in self.text_input.text.split("\n"):
                    doc.add_paragraph(paragraph_text)
                doc.save(self.current_file)
            elif file_ext == ".xlsx":
                wb = openpyxl.Workbook()
                ws = wb.active
                for row in self.text_input.text.split("\n"):
                    row_cells = row.split("\t")
                    ws.append(row_cells)
                wb.save(self.current_file)
            else:
                with open(self.current_file, 'w', encoding='utf-8') as f:
                    f.write(self.text_input.text)
            instance.parent.parent.parent.dismiss()
if __name__ == "__main__":
    TextEditor().run()

