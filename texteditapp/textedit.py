import os
import docx
import openpyxl
from kivy.lang import Builder
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.popup import Popup
from kivy.properties import ObjectProperty
from kivy.app import App
from kivy.uix.button import Button
from kivy.uix.filechooser import FileChooserIconView

KV = '''
BoxLayout:
    orientation: "vertical"

    BoxLayout:
        size_hint_y: None
        height: "48dp"

        Button:
            text: "Open"
            on_release: app.open_file()

        Button:
            text: "Save"
            on_release: app.save_file()

        Button:
            text: "Save As"
            on_release: app.save_file_as()

    TextInput:
        id: text_input
        focus: True
        font_size: '16sp'
        on_parent: app.text_input = self  # Add this line
'''

class TextEditor(App):
    text_input = ObjectProperty(None)
    current_file = ''

    def build(self):
        return Builder.load_string(KV)

    def open_file(self):
        filechooser = FileChooserIconView(path=os.getcwd())
        popup = Popup(title="Open File", content=filechooser, size_hint=(0.8, 0.8))
        filechooser.bind(on_submit=self.load_file)
        popup.open()

    def load_file(self, instance, filepaths, *args):
        if filepaths:
            self.current_file = filepaths[0]
            file_ext = os.path.splitext(self.current_file)[1].lower()

            if file_ext == ".docx":
                try:
                    doc = docx.Document(self.current_file)
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text is not None:
                            full_text.append(paragraph.text)

                    self.text_input.text = "\n".join(full_text)
                except Exception as e:
                    print(f"Error: {e}")

            elif file_ext == ".xlsx":
                try:
                    wb = openpyxl.load_workbook(self.current_file)
                    sheet = wb.active
                    full_text = []
                    for row in sheet.iter_rows():
                        row_text = " ".join([str(cell.value) for cell in row])
                        full_text.append(row_text)
                    self.text_input.text = "\n".join(full_text)
                except Exception as e:
                    print(f"Error: {e}")

            else:
                try:
                    with open(self.current_file, 'r', encoding='utf-8') as f:
                        self.text_input.text = f.read()
                except UnicodeDecodeError as e:
                    print(f"Error: {e}")

            instance.parent.parent.parent.dismiss()

    def save_file(self):
        if self.current_file:
            file_ext = os.path.splitext(self.current_file)[1].lower()
            if file_ext == ".docx":
                doc = docx.Document()
                for paragraph_text in self.text_input.text.split("\n"):
                    doc.add_paragraph(paragraph_text)
                doc.save(self.current_file)
            else:
                with open(self.current_file, 'w', encoding='utf-8') as f:
                    f.write(self.text_input.text)
        else:
            self.save_file_as()
            
    def save_file_as(self):
        filechooser = FileChooserIconView(path=os.getcwd(), dirselect=True)
        popup = Popup(title="Save File As", content=filechooser, size_hint=(0.8, 0.8))
        filechooser.bind(on_submit=self.save_new_file_with_popup)
        popup.open()

    def save_new_file_with_popup(self, instance, filepaths, *args):
        if filepaths:
            filename = os.path.basename(instance.selection[0])
            self.current_file = os.path.join(filepaths[0], filename)
            file_ext = os.path.splitext(self.current_file)[1].lower()
            if file_ext == ".docx":
                doc = docx.Document()
                for paragraph_text in self.text_input.text.split("\n"):
                    doc.add_paragraph(paragraph_text)
                doc.save(self.current_file)
            elif file_ext == ".xlsx":
                wb = openpyxl.Workbook()
                sheet = wb.active
                for row_text in self.text_input.text.split("\n"):
                    row_values = row_text.split()
                    sheet.append(row_values)
                wb.save(self.current_file)
            else:
                with open(self.current_file, 'w', encoding='utf-8') as f:
                    f.write(self.text_input.text)
            instance.parent.parent.parent.dismiss()

if __name__ == "__main__":
    TextEditor().run()
            
    def save_file_as(self):
        filechooser = FileChooserIconView(path=os.getcwd(), dirselect=True)
        popup = Popup(title="Save File As", content=filechooser, size_hint=(0.8, 0.8))
        filechooser.bind(on_submit=self.save_new_file_with_popup)
        popup.open()

    def save_new_file_with_popup(self, instance, filepaths, *args):
        if filepaths:
            filename = os.path.basename(instance.selection[0])
            self.current_file = os.path.join(filepaths[0], filename)
            file_ext = os.path.splitext(self.current_file)[1].lower()
            if file_ext == ".docx":
                doc = docx.Document()
                for paragraph_text in self.text_input.text.split("\n"):
                    doc.add_paragraph(paragraph_text)
                doc.save(self.current_file)
            else:
                with open(self.current_file, 'w', encoding='utf-8') as f:
                    f.write(self.text_input.text)
            instance.parent.parent.parent.dismiss()

if __name__ == "__main__":
    TextEditor().run()
